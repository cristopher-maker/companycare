import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ==========================================
# 1. GENERATE WORD DOCUMENT (.DOCX)
# ==========================================
doc = docx.Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

COLOR_NAVY = RGBColor(18, 60, 79)       # #123C4F
COLOR_ORANGE = RGBColor(242, 122, 94)   # #F27A5E
COLOR_DARK = RGBColor(44, 62, 80)       # #2C3E50
COLOR_MUTED = RGBColor(107, 123, 133)   # #6B7B85

def set_cell_bg(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# Title Header
p_title = doc.add_paragraph()
r_brand = p_title.add_run("CompanyCare | ")
r_brand.font.size = Pt(22)
r_brand.font.bold = True
r_brand.font.color.rgb = COLOR_NAVY

r_t = p_title.add_run("Ficha Técnica Corporativa")
r_t.font.size = Pt(20)
r_t.font.bold = True
r_t.font.color.rgb = COLOR_ORANGE

p_sub = doc.add_paragraph()
r_s = p_sub.add_run("Solución Integral de Bienestar, Asistencia y Cuidado Senior para Colaboradores")
r_s.font.size = Pt(11)
r_s.font.italic = True
r_s.font.color.rgb = COLOR_MUTED
p_sub.paragraph_format.space_after = Pt(16)

# Section 1: Executive Summary
h1 = doc.add_paragraph()
r_h1 = h1.add_run("1. Resumen Ejecutivo")
r_h1.font.size = Pt(14)
r_h1.font.bold = True
r_h1.font.color.rgb = COLOR_NAVY

p_body = doc.add_paragraph()
p_body.paragraph_format.space_after = Pt(12)
r_b = p_body.add_run(
    "CompanyCare es el programa corporativo integral diseñado para acompañar y apoyar a los colaboradores que "
    "asumen responsabilidades de cuidado de adultos mayores y familiares dependientes. Conectamos a los empleados con "
    "expertos profesionales (gerontólogos y trabajadores sociales), una red verificada de proveedores de salud y residencias, "
    "guías operativas y herramientas digitales que reducen el estrés familiar y mejoran directamente la productividad laboral."
)
r_b.font.size = Pt(10)
r_b.font.color.rgb = COLOR_DARK

# KPI Cards Table
t_kpi = doc.add_table(rows=1, cols=3)
t_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
kpi_data = [
    ("-18%", "Reducción de Ausentismo", "Disminución de licencias e interrupciones laborales por emergencias de cuidado."),
    ("4.8 / 5", "Satisfacción Empleados", "Evaluación promedio del acompañamiento 1 a 1 de Care Experts."),
    ("72 Horas", "Tiempo de Respuesta", "Resolución y orientación completa desde la primera consulta hasta el plan concreto.")
]

for i, (num, label, desc) in enumerate(kpi_data):
    cell = t_kpi.rows[0].cells[i]
    cell.width = Inches(2.2)
    set_cell_bg(cell, "FFF0EB")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    rn = p.add_run(f"{num}\n")
    rn.font.size = Pt(18)
    rn.font.bold = True
    rn.font.color.rgb = COLOR_ORANGE
    
    rl = p.add_run(f"{label}\n")
    rl.font.size = Pt(9.5)
    rl.font.bold = True
    rl.font.color.rgb = COLOR_NAVY
    
    rd = p.add_run(desc)
    rd.font.size = Pt(8.5)
    rd.font.color.rgb = COLOR_MUTED

doc.add_paragraph().paragraph_format.space_after = Pt(16)

# Section 2: Services & Modules
h2 = doc.add_paragraph()
r_h2 = h2.add_run("2. Pilares de Servicios Incluidos")
r_h2.font.size = Pt(14)
r_h2.font.bold = True
r_h2.font.color.rgb = COLOR_NAVY

modules = [
    ("Asesoría con Care Experts", "Atención personal 1 a 1 mediante chat y videollamada con gerontólogos, trabajadores sociales y gestores de casos."),
    ("Red de Proveedores Verificados", "Acceso directo y evaluado a residencias de larga estadía, cuidadores domiciliarios, enfermería y servicios médicos especializados."),
    ("Biblioteca de Recursos & Guías", "Contenido descargable paso a paso, checklists operativos y opciones de financiamiento y trámites legales senior."),
    ("Formación y Bienestar Emocional", "Talleres online y webinars para reducir el síndrome de burnout del cuidador y mejorar el bienestar del empleado."),
    ("Seguimiento de Casos en Tiempo Real", "Plataforma digital con expediente completo, estado transparente y acompañamiento continuo de cada solicitud."),
    ("Panel Empresa (RR.HH. Analytics)", "Dashboard ejecutivo para Gestión de Personas con métricas de adopción, co-branding y reportes de impacto anónimos.")
]

t_mod = doc.add_table(rows=1, cols=2)
t_mod.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t_mod.rows[0].cells
hdr[0].width = Inches(2.3)
hdr[1].width = Inches(4.4)
set_cell_bg(hdr[0], "123C4F")
set_cell_bg(hdr[1], "123C4F")

r_m0 = hdr[0].paragraphs[0].add_run("Módulo / Servicio")
r_m0.font.bold = True
r_m0.font.color.rgb = RGBColor(255, 255, 255)
r_m0.font.size = Pt(9.5)

r_m1 = hdr[1].paragraphs[0].add_run("Descripción y Alcance Corporativo")
r_m1.font.bold = True
r_m1.font.color.rgb = RGBColor(255, 255, 255)
r_m1.font.size = Pt(9.5)

for name, dsc in modules:
    row_c = t_mod.add_row().cells
    row_c[0].width = Inches(2.3)
    row_c[1].width = Inches(4.4)
    
    p0 = row_c[0].paragraphs[0]
    rn = p0.add_run(name)
    rn.font.bold = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = COLOR_NAVY
    
    p1 = row_c[1].paragraphs[0]
    rd = p1.add_run(dsc)
    rd.font.size = Pt(9)
    rd.font.color.rgb = COLOR_DARK

doc.add_paragraph().paragraph_format.space_after = Pt(16)

# Section 3: Plans Comparison
h3 = doc.add_paragraph()
r_h3 = h3.add_run("3. Cobertura de Planes Corporativos")
r_h3.font.size = Pt(14)
r_h3.font.bold = True
r_h3.font.color.rgb = COLOR_NAVY

t_plan = doc.add_table(rows=1, cols=4)
t_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
phdrs = ["Característica", "Plan Starter (1-100)", "Plan Corporate (101-500)", "Plan Enterprise (500+)"]
pw = [Inches(2.2), Inches(1.5), Inches(1.5), Inches(1.5)]

for j, title in enumerate(phdrs):
    c = t_plan.rows[0].cells[j]
    c.width = pw[j]
    set_cell_bg(c, "F27A5E")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(9)

prows = [
    ("Atención Care Experts (Chat/Video)", "Incluido", "Incluido", "Ilimitado + Gestor Dedicado"),
    ("Red de Proveedores Verificados", "Acceso Catálogo", "Tarifas Preferenciales", "Red a Medida + Descuentos"),
    ("Biblioteca y Guías", "Acceso Completo", "Acceso Completo", "Personalizado con Co-Branding"),
    ("Talleres y Webinars", "2 al año", "Frecuencia Mensual", "Programa a Medida + Grabados"),
    ("Dashboard RR.HH. Analítica", "Básico", "Avanzado", "Custom / Integración API"),
    ("Onboarding y Kits RR.HH.", "Estándar", "Personalizado", "Llave en mano")
]

for r_data in prows:
    rc = t_plan.add_row().cells
    for k, val in enumerate(r_data):
        rc[k].width = pw[k]
        p = rc[k].paragraphs[0]
        if k > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        r.font.size = Pt(8.5)
        if k == 0:
            r.font.bold = True
            r.font.color.rgb = COLOR_NAVY
        else:
            r.font.color.rgb = COLOR_DARK

doc.save("c:/Users/crist/OneDrive/Documentos/Company Care/Ficha_Tecnica_CompanyCare.docx")
print("Word generated successfully.")

# ==========================================
# 2. GENERATE EXCEL DOCUMENT (.XLSX)
# ==========================================
wb = openpyxl.Workbook()

# Sheet 1: Resumen y Módulos
ws1 = wb.active
ws1.title = "Ficha Técnica y Servicios"

ws1.views.sheetView[0].showGridLines = True

font_title = Font(name="Segoe UI", size=16, bold=True, color="123C4F")
font_sub = Font(name="Segoe UI", size=11, italic=True, color="6B7B85")
font_h1 = Font(name="Segoe UI", size=12, bold=True, color="FFFFFF")
font_bold = Font(name="Segoe UI", size=10, bold=True, color="123C4F")
font_regular = Font(name="Segoe UI", size=10, color="2C3E50")
font_kpi_num = Font(name="Segoe UI", size=16, bold=True, color="F27A5E")

fill_navy = PatternFill(start_color="123C4F", end_color="123C4F", fill_type="solid")
fill_orange = PatternFill(start_color="F27A5E", end_color="F27A5E", fill_type="solid")
fill_soft = PatternFill(start_color="FFF0EB", end_color="FFF0EB", fill_type="solid")

thin_border = Border(
    left=Side(style='thin', color='E5E9EB'),
    right=Side(style='thin', color='E5E9EB'),
    top=Side(style='thin', color='E5E9EB'),
    bottom=Side(style='thin', color='E5E9EB')
)

ws1["A1"] = "CompanyCare - Ficha Técnica Corporativa"
ws1["A1"].font = font_title
ws1["A2"] = "Programa Corporativo de Bienestar y Cuidado Senior para Colaboradores"
ws1["A2"].font = font_sub

# Metrics Summary Table
ws1["A4"] = "Indicador de Impacto"
ws1["B4"] = "Valor"
ws1["C4"] = "Detalle"
for col in ["A4", "B4", "C4"]:
    ws1[col].font = font_h1
    ws1[col].fill = fill_orange
    ws1[col].alignment = Alignment(horizontal="center", vertical="center")

kpis_xl = [
    ("Reducción de Ausentismo", "-18%", "Menos licencias e interrupciones por emergencias de cuidado de familiares."),
    ("Satisfacción de Empleados", "4.8 / 5", "Calificación promedio del acompañamiento 1 a 1 de los Care Experts."),
    ("Tiempo de Respuesta", "72 Horas", "De la primera consulta a un plan de acción concreto.")
]

for row_idx, (ind, val, det) in enumerate(kpis_xl, start=5):
    ws1[f"A{row_idx}"] = ind
    ws1[f"A{row_idx}"].font = font_bold
    ws1[f"B{row_idx}"] = val
    ws1[f"B{row_idx}"].font = font_kpi_num
    ws1[f"B{row_idx}"].alignment = Alignment(horizontal="center")
    ws1[f"C{row_idx}"] = det
    ws1[f"C{row_idx}"].font = font_regular
    
    for c in ["A", "B", "C"]:
        ws1[f"{c}{row_idx}"].border = thin_border
        ws1[f"{c}{row_idx}"].fill = fill_soft

# Modules Table
ws1["A9"] = "Módulo de Servicio"
ws1["B9"] = "Público Objetivo"
ws1["C9"] = "Descripción y Beneficio Corporativo"
for col in ["A9", "B9", "C9"]:
    ws1[col].font = font_h1
    ws1[col].fill = fill_navy
    ws1[col].alignment = Alignment(horizontal="center", vertical="center")

mods_xl = [
    ("Asesoría con Care Experts", "Empleado / Familia", "Orientación 1 a 1 por chat y videollamada con gerontólogos y trabajadores sociales."),
    ("Red de Proveedores Verificados", "Empleado / Familia", "Acceso a residencias, cuidadores domiciliarios y salud con datos y evaluaciones reales."),
    ("Biblioteca de Recursos & Guías", "Empleado / Familia", "Guías paso a paso, checklists operativas y opciones de financiamiento senior."),
    ("Formación y Bienestar Emocional", "Empleado", "Cursos online y talleres de manejo del estrés para colaboradores cuidadores."),
    ("Seguimiento Digital de Casos", "Empleado / Care Expert", "Expediente digitalizado con estado en tiempo real y acompañamiento continuo."),
    ("Panel Empresa (RR.HH.)", "Recursos Humanos", "Métricas de adopción, co-branding, onboarding personalizado e indicadores de impacto.")
]

for r_idx, (m_name, m_aud, m_desc) in enumerate(mods_xl, start=10):
    ws1[f"A{r_idx}"] = m_name
    ws1[f"A{r_idx}"].font = font_bold
    ws1[f"B{r_idx}"] = m_aud
    ws1[f"B{r_idx}"].font = font_regular
    ws1[f"B{r_idx}"].alignment = Alignment(horizontal="center")
    ws1[f"C{r_idx}"] = m_desc
    ws1[f"C{r_idx}"].font = font_regular
    
    for c in ["A", "B", "C"]:
        ws1[f"{c}{r_idx}"].border = thin_border

# Sheet 2: Comparativo de Planes
ws2 = wb.create_sheet(title="Planes y Cobertura")
ws2.views.sheetView[0].showGridLines = True

ws2["A1"] = "Comparativo de Planes Corporativos CompanyCare"
ws2["A1"].font = font_title

headers_p = ["Característica", "Starter (1-100 emp)", "Corporate (101-500 emp)", "Enterprise (500+ emp)"]
for c_idx, h_text in enumerate(headers_p, start=1):
    cell = ws2.cell(row=3, column=c_idx, value=h_text)
    cell.font = font_h1
    cell.fill = fill_orange
    cell.alignment = Alignment(horizontal="center", vertical="center")

plans_xl = [
    ("Atención Care Experts (Chat/Video)", "Incluido", "Incluido", "Ilimitado + Gestor Dedicado"),
    ("Red de Proveedores Verificados", "Acceso Catálogo", "Tarifas Preferenciales", "Red a Medida + Descuentos"),
    ("Biblioteca y Guías", "Acceso Completo", "Acceso Completo", "Personalizado con Co-Branding"),
    ("Talleres y Webinars", "2 al año", "Frecuencia Mensual", "Programa a Medida + Grabados"),
    ("Dashboard RR.HH. Analítica", "Básico", "Avanzado", "Custom / Integración API"),
    ("Onboarding y Kits RR.HH.", "Estándar", "Personalizado", "Llave en mano")
]

for r_i, row_t in enumerate(plans_xl, start=4):
    for c_i, val_t in enumerate(row_t, start=1):
        cell = ws2.cell(row=r_i, column=c_i, value=val_t)
        cell.font = font_bold if c_i == 1 else font_regular
        cell.border = thin_border
        if c_i > 1:
            cell.alignment = Alignment(horizontal="center")

# Auto-adjust column widths
for sheet in [ws1, ws2]:
    for col in sheet.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        sheet.column_dimensions[col_letter].width = min(max(max_len + 4, 12), 65)

wb.save("c:/Users/crist/OneDrive/Documentos/Company Care/Desglose_Servicios_CompanyCare.xlsx")
print("Excel generated successfully.")
