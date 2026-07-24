import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ==========================================
# 1. GENERATE WORD (.DOCX)
# ==========================================
doc = docx.Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

COLOR_NAVY = RGBColor(18, 60, 79)       # #123C4F
COLOR_ORANGE = RGBColor(242, 122, 94)   # #F27A5E
COLOR_DARK = RGBColor(44, 62, 80)       # #2C3E50
COLOR_MUTED = RGBColor(107, 123, 133)   # #6B7B85

def set_cell_bg(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_sec_title(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COLOR_NAVY
    return p

def add_table_fields(fields_list):
    t = doc.add_table(rows=len(fields_list), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (lbl, val) in enumerate(fields_list):
        c1, c2 = t.rows[idx].cells[0], t.rows[idx].cells[1]
        c1.width = Inches(2.6)
        c2.width = Inches(4.2)
        
        set_cell_bg(c1, "F4F7F9")
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(lbl)
        r1.font.size = Pt(9.5)
        r1.font.bold = True
        r1.font.color.rgb = COLOR_NAVY
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(val if val else "__________________________________________________")
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = COLOR_MUTED if not val else COLOR_DARK

# Header
p_t = doc.add_paragraph()
r_b = p_t.add_run("CompanyCare | ")
r_b.font.size = Pt(20)
r_b.font.bold = True
r_b.font.color.rgb = COLOR_NAVY

r_t = p_t.add_run("Ficha de Ingreso para Empleados")
r_t.font.size = Pt(18)
r_t.font.bold = True
r_t.font.color.rgb = COLOR_ORANGE

p_sub = doc.add_paragraph()
r_s = p_sub.add_run("Formulario Oficial de Diagnóstico y Registro de Caso de Cuidado")
r_s.font.size = Pt(10.5)
r_s.font.italic = True
r_s.font.color.rgb = COLOR_MUTED
p_sub.paragraph_format.space_after = Pt(8)

# 1. Datos Del Adulto Mayor
add_sec_title("1. Datos Del Adulto Mayor")
add_table_fields([
    ("Nombre completo:", ""),
    ("RUT / Documento:", ""),
    ("Fecha de nacimiento:", "DD / MM / AAAA"),
    ("Edad:", ""),
    ("Sexo:", ""),
    ("Nacionalidad:", ""),
    ("Dirección:", ""),
    ("Comuna:", ""),
    ("Teléfono:", ""),
    ("Correo, si aplica:", "")
])

# 2. Contacto Responsable
add_sec_title("2. Contacto Responsable")
add_table_fields([
    ("Nombre completo:", ""),
    ("Relación con el adulto mayor:", ""),
    ("RUT / Documento:", ""),
    ("Teléfono:", ""),
    ("Correo:", ""),
    ("Dirección:", ""),
    ("¿Es cuidador principal?:", "[  ] Sí    [  ] No")
])

# 3. Información De Salud
add_sec_title("3. Información De Salud")
add_table_fields([
    ("Previsión de salud:", "[  ] Fonasa    [  ] Isapre    [  ] Particular    [  ] Otra"),
    ("Diagnósticos principales:", ""),
    ("Enfermedades crónicas:", ""),
    ("Medicamentos actuales:", ""),
    ("Alergias:", ""),
    ("Médico tratante:", ""),
    ("Teléfono del médico / centro:", ""),
    ("Última hospitalización, si aplica:", ""),
    ("Observaciones médicas relevantes:", "")
])

# 4. Nivel De Dependencia
add_sec_title("4. Nivel De Dependencia")
add_table_fields([
    ("Se moviliza solo/a:", "[  ] Sí    [  ] No    [  ] Con ayuda"),
    ("Usa ayudas técnicas:", "[  ] Bastón    [  ] Andador    [  ] Silla de ruedas    [  ] Otro"),
    ("Requiere ayuda para bañarse:", "[  ] Sí    [  ] No"),
    ("Requiere ayuda para vestirse:", "[  ] Sí    [  ] No"),
    ("Requiere ayuda para alimentarse:", "[  ] Sí    [  ] No"),
    ("Requiere ayuda para tomar medicamentos:", "[  ] Sí    [  ] No"),
    ("Requiere supervisión permanente:", "[  ] Sí    [  ] No"),
    ("Riesgo de caídas:", "[  ] Bajo    [  ] Medio    [  ] Alto")
])

# 5. Estado Cognitivo Y Emocional
add_sec_title("5. Estado Cognitivo Y Emocional")
add_table_fields([
    ("Orientado/a en tiempo y espacio:", "[  ] Sí    [  ] No    [  ] Parcial"),
    ("Diagnóstico cognitivo:", "[  ] Ninguno    [  ] Deterioro cognitivo    [  ] Alzheimer    [  ] Demencia    [  ] Otro"),
    ("Cambios de conducta:", ""),
    ("Estado de ánimo frecuente:", ""),
    ("Ansiedad, tristeza o aislamiento:", ""),
    ("Observaciones:", "")
])

# 6. Situación Familiar Y De Cuidado
add_sec_title("6. Situación Familiar Y De Cuidado")
add_table_fields([
    ("Vive solo/a:", "[  ] Sí    [  ] No"),
    ("Vive con:", ""),
    ("Cuidador principal:", ""),
    ("Horarios en que requiere apoyo:", ""),
    ("Red familiar disponible:", ""),
    ("Principales preocupaciones de la familia:", ""),
    ("Urgencia del caso:", "[  ] Baja    [  ] Media    [  ] Alta")
])

# 7. Servicio Solicitado
add_sec_title("7. Servicio Solicitado")
add_table_fields([
    ("Servicios requeridos (Marcar con X):", 
     "[  ] Asesoría de cuidado\n"
     "[  ] Evaluación del caso\n"
     "[  ] Búsqueda de cuidador/a\n"
     "[  ] Residencia para adulto mayor\n"
     "[  ] Apoyo médico o enfermería\n"
     "[  ] Apoyo psicológico\n"
     "[  ] Orientación legal / financiera\n"
     "[  ] Otro: ___________________________")
])

# 8. Descripción Del Caso
add_sec_title("8. Descripción Del Caso")
p_desc = doc.add_paragraph()
r_d = p_desc.add_run("Describe brevemente la situación actual del adulto mayor, qué problema necesitan resolver y qué tipo de apoyo esperan recibir:\n\n")
r_d.font.size = Pt(9.5)
r_d.font.bold = True
r_d.font.color.rgb = COLOR_NAVY

r_lines = p_desc.add_run(
    "____________________________________________________________________________________________________\n"
    "____________________________________________________________________________________________________\n"
    "____________________________________________________________________________________________________\n"
)
r_lines.font.size = Pt(9)
r_lines.font.color.rgb = COLOR_MUTED

doc.save("c:/Users/crist/OneDrive/Documentos/Company Care/Ficha_CompanyCare_Empleados.docx")
print("Word employee intake generated.")

# ==========================================
# 2. GENERATE EXCEL (.XLSX)
# ==========================================
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Ficha Company Care"
ws.views.sheetView[0].showGridLines = True

font_t = Font(name="Segoe UI", size=15, bold=True, color="123C4F")
font_s = Font(name="Segoe UI", size=10, italic=True, color="6B7B85")
font_sec = Font(name="Segoe UI", size=11, bold=True, color="FFFFFF")
font_lbl = Font(name="Segoe UI", size=9.5, bold=True, color="123C4F")
font_val = Font(name="Segoe UI", size=9.5, color="2C3E50")

fill_sec = PatternFill(start_color="123C4F", end_color="123C4F", fill_type="solid")
fill_lbl = PatternFill(start_color="F4F7F9", end_color="F4F7F9", fill_type="solid")

thin_b = Border(
    left=Side(style='thin', color='E5E9EB'),
    right=Side(style='thin', color='E5E9EB'),
    top=Side(style='thin', color='E5E9EB'),
    bottom=Side(style='thin', color='E5E9EB')
)

ws["A1"] = "Ficha Company Care"
ws["A1"].font = font_t
ws["A2"] = "Formulario Oficial de Diagnóstico para Empleados"
ws["A2"].font = font_s

sections_xl = [
    ("1. Datos Del Adulto Mayor", [
        ("Nombre completo:", ""), ("RUT / Documento:", ""), ("Fecha de nacimiento:", "DD/MM/AAAA"),
        ("Edad:", ""), ("Sexo:", ""), ("Nacionalidad:", ""), ("Dirección:", ""),
        ("Comuna:", ""), ("Teléfono:", ""), ("Correo, si aplica:", "")
    ]),
    ("2. Contacto Responsable", [
        ("Nombre completo:", ""), ("Relación con el adulto mayor:", ""), ("RUT / Documento:", ""),
        ("Teléfono:", ""), ("Correo:", ""), ("Dirección:", ""), ("¿Es cuidador principal?:", "Sí / No")
    ]),
    ("3. Información De Salud", [
        ("Previsión de salud:", "Fonasa / Isapre / Particular / Otra"), ("Diagnósticos principales:", ""),
        ("Enfermedades crónicas:", ""), ("Medicamentos actuales:", ""), ("Alergias:", ""),
        ("Médico tratante:", ""), ("Teléfono del médico o centro de salud:", ""),
        ("Última hospitalización, si aplica:", ""), ("Observaciones médicas relevantes:", "")
    ]),
    ("4. Nivel De Dependencia", [
        ("Se moviliza solo/a:", "Sí / No / Con ayuda"), ("Usa ayudas técnicas:", "Bastón / Andador / Silla de ruedas / Otro"),
        ("Requiere ayuda para bañarse:", "Sí / No"), ("Requiere ayuda para vestirse:", "Sí / No"),
        ("Requiere ayuda para alimentarse:", "Sí / No"), ("Requiere ayuda para tomar medicamentos:", "Sí / No"),
        ("Requiere supervisión permanente:", "Sí / No"), ("Riesgo de caídas:", "Bajo / Medio / Alto")
    ]),
    ("5. Estado Cognitivo Y Emocional", [
        ("Orientado/a en tiempo y espacio:", "Sí / No / Parcial"),
        ("Diagnóstico cognitivo:", "Ninguno / Deterioro cognitivo / Alzheimer / Demencia / Otro"),
        ("Cambios de conducta:", ""), ("Estado de ánimo frecuente:", ""),
        ("Ansiedad, tristeza o aislamiento:", ""), ("Observaciones:", "")
    ]),
    ("6. Situación Familiar Y De Cuidado", [
        ("Vive solo/a:", "Sí / No"), ("Vive con:", ""), ("Cuidador principal:", ""),
        ("Horarios en que requiere apoyo:", ""), ("Red familiar disponible:", ""),
        ("Principales preocupaciones de la familia:", ""), ("Urgencia del caso:", "Baja / Media / Alta")
    ]),
    ("7. Servicio Solicitado", [
        ("Servicios requeridos:", "Asesoría / Evaluación / Búsqueda de cuidador / Residencia / Apoyo médico / Psicológico / Legal / Otro")
    ]),
    ("8. Descripción Del Caso", [
        ("Situación actual y apoyo esperado:", "Describe brevemente la situación actual del adulto mayor, qué problema necesitan resolver y qué apoyo esperan recibir...")
    ])
]

curr_row = 4
for stitle, fields in sections_xl:
    ws.cell(row=curr_row, column=1, value=stitle)
    ws.cell(row=curr_row, column=2, value="")
    ws.merge_cells(start_row=curr_row, start_column=1, end_row=curr_row, end_column=2)
    
    cell_h = ws.cell(row=curr_row, column=1)
    cell_h.font = font_sec
    cell_h.fill = fill_sec
    cell_h.alignment = Alignment(horizontal="left", vertical="center")
    curr_row += 1
    
    for flbl, fval in fields:
        c1 = ws.cell(row=curr_row, column=1, value=flbl)
        c2 = ws.cell(row=curr_row, column=2, value=fval)
        
        c1.font = font_lbl
        c1.fill = fill_lbl
        c1.border = thin_b
        
        c2.font = font_val
        c2.border = thin_b
        curr_row += 1
    curr_row += 1

ws.column_dimensions["A"].width = 36
ws.column_dimensions["B"].width = 65

wb.save("c:/Users/crist/OneDrive/Documentos/Company Care/Ficha_CompanyCare_Empleados.xlsx")
print("Excel employee intake generated.")
