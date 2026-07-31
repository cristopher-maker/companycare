import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Page Setup
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Color Palette
COLOR_NAVY = RGBColor(18, 60, 79)       # #123C4F
COLOR_CYAN = RGBColor(15, 110, 86)      # #0F6E56
COLOR_DARK = RGBColor(4, 36, 51)        # #042433
COLOR_TEXT = RGBColor(40, 44, 52)       # Dark Gray

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("COMPANY CARE\nInforme de Funcionalidades del Sitio y Anexo Piloto")
title_run.font.name = 'Calibri'
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = COLOR_NAVY
title_p.paragraph_format.space_after = Pt(4)

# Subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("Plataforma Corporativa de Cuidado Senior | Operado por Comercial Ovni")
sub_run.font.name = 'Calibri'
sub_run.font.size = Pt(12)
sub_run.font.italic = True
sub_run.font.color.rgb = COLOR_CYAN
sub_p.paragraph_format.space_after = Pt(24)

# Horizontal Divider Line
p_hr = doc.add_paragraph()
p_hr.paragraph_format.space_after = Pt(18)
p_hr_run = p_hr.add_run("_________________________________________________________________________________")
p_hr_run.font.color.rgb = RGBColor(200, 210, 220)

# Section 1: Intro
h1 = doc.add_heading("1. Descripción General de la Plataforma", level=1)
h1.runs[0].font.color.rgb = COLOR_NAVY
h1.runs[0].font.size = Pt(16)
h1.runs[0].font.bold = True

p_intro = doc.add_paragraph(
    "Company Care es la solución digital corporativa orientada a conectar a las empresas y sus colaboradores "
    "con una red integral de expertos en gerontología, proveedores auditados y recursos para la gestión y "
    "acompañamiento del cuidado de adultos mayores, reduciendo el estrés del trabajador y el ausentismo laboral."
)
p_intro.runs[0].font.name = 'Calibri'
p_intro.runs[0].font.size = Pt(11)
p_intro.runs[0].font.color.rgb = COLOR_TEXT
p_intro.paragraph_format.space_after = Pt(14)

# Section 2: Funcionalidades del Sitio Web
h2 = doc.add_heading("2. Funcionalidades del Sitio Web y Plataforma", level=1)
h2.runs[0].font.color.rgb = COLOR_NAVY
h2.runs[0].font.size = Pt(16)
h2.runs[0].font.bold = True

features_data = [
    ("Portal del Colaborador (Empleado)", [
        "Acceso Corporativo Seguro: Inicio de sesión mediante correo institucional de la empresa.",
        "Asesoría con Care Experts: Chat y videollamada con gerontólogos y trabajadores sociales para evaluar cada caso.",
        "Red de Proveedores Verificados: Catálogo filtrable de residencias auditadas, cuidadores a domicilio y especialistas médicos en Chile.",
        "Biblioteca de Recursos: Guías paso a paso descargables, checklists gerontológicos y orientación legal/financiera.",
        "Módulo de Bienestar y Formación: Cursos y cápsulas online para prevenir el agotamiento (burnout) del cuidador.",
        "Seguimiento de Casos: Monitoreo en tiempo real del estado de cada solicitud."
    ]),
    ("Panel de Recursos Humanos (Dashboard HR)", [
        "Métricas de Impacto en Tiempo Real: Visualización de ausentismo reducido, nivel de adopción del beneficio y satisfacción (4.8/5).",
        "Personalización Co-Branding: Adaptación de la plataforma con la identidad corporativa de la empresa cliente.",
        "Gestión de Accesos al Beneficio: Control de habilitación y asignación de colaboradores activos.",
        "Reportes Ejecutivos: Generación de reportes descargables para la dirección general de la empresa."
    ]),
    ("Sitio Web y Canales Públicos", [
        "Sitio Informativo e Interactivo: Muestra detallada de servicios, modelo de atención e indicadores clave.",
        "Cotizador y Formulario Corporativo: Solicitud directa de propuestas y demostraciones personalizadas."
    ])
]

for category, items in features_data:
    p_cat = doc.add_paragraph()
    run_cat = p_cat.add_run(f"• {category}")
    run_cat.font.name = 'Calibri'
    run_cat.font.size = Pt(13)
    run_cat.font.bold = True
    run_cat.font.color.rgb = COLOR_DARK
    p_cat.paragraph_format.space_before = Pt(8)
    p_cat.paragraph_format.space_after = Pt(4)
    
    for item in items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.left_indent = Inches(0.3)
        p_item.paragraph_format.space_after = Pt(3)
        run_item = p_item.add_run(f"- {item}")
        run_item.font.name = 'Calibri'
        run_item.font.size = Pt(10.5)
        run_item.font.color.rgb = COLOR_TEXT

# Section 3: Anexo Piloto
h3 = doc.add_heading("3. Anexo de Condiciones de Asistencia Técnico-Funcional (Piloto)", level=1)
h3.runs[0].font.color.rgb = COLOR_NAVY
h3.runs[0].font.size = Pt(16)
h3.runs[0].font.bold = True

p_anexo_desc = doc.add_paragraph(
    "El presente Anexo regula las condiciones de asistencia técnico-funcional provistas por Comercial Ovni "
    "para la plataforma Company Care durante el período piloto:"
)
p_anexo_desc.runs[0].font.name = 'Calibri'
p_anexo_desc.runs[0].font.size = Pt(11)
p_anexo_desc.runs[0].font.color.rgb = COLOR_TEXT
p_anexo_desc.paragraph_format.space_after = Pt(12)

# Table setup
table = doc.add_table(rows=5, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

headers = ["Concepto", "Detalle y Condición"]
hdr_cells = table.rows[0].cells
for i, title in enumerate(headers):
    hdr_cells[i].text = title
    set_cell_background(hdr_cells[i], "123C4F")
    set_cell_margins(hdr_cells[i], top=160, bottom=160, left=180, right=180)
    for p in hdr_cells[i].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

table_data = [
    ("Canales Habilitados", "Soporte vía correo electrónico a contacto@comercialovni.cl y chat de mensajería al número corporativo coordinado con la contraparte del Cliente."),
    ("Horario de Atención", "Lunes a viernes de 09:00 a 18:00 horas (Hora de Chile Continental). Se excluyen días sábados, domingos y festivos en Chile."),
    ("Tiempos de Respuesta", "Máximo 24 horas hábiles desde la recepción formal del requerimiento para entregar el acuse de recibo o estado técnico inicial."),
    ("Exclusiones", "El soporte no incluye reparación de hardware, configuraciones de conectividad interna del Cliente, ni desarrollos de software a medida ajenos al alcance original del piloto.")
]

for row_idx, (concept, detail) in enumerate(table_data, start=1):
    row_cells = table.rows[row_idx].cells
    row_cells[0].text = concept
    row_cells[1].text = detail
    
    bg_color = "F4F7F9" if row_idx % 2 == 1 else "FFFFFF"
    
    for col_idx in range(2):
        set_cell_background(row_cells[col_idx], bg_color)
        set_cell_margins(row_cells[col_idx], top=140, bottom=140, left=180, right=180)
        for p in row_cells[col_idx].paragraphs:
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.color.rgb = COLOR_TEXT
                if col_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = COLOR_DARK

# Column widths
widths = [Inches(2.2), Inches(4.3)]
for row in table.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w

# Footer note
p_foot = doc.add_paragraph()
p_foot.paragraph_format.space_before = Pt(28)
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_foot = p_foot.add_run("Documento Oficial - Company Care | Comercial Ovni 2026")
r_foot.font.name = 'Calibri'
r_foot.font.size = Pt(9)
r_foot.font.italic = True
r_foot.font.color.rgb = RGBColor(140, 150, 160)

output_path = "c:/Users/crist/OneDrive/Documentos/Company Care/Funcionalidades_y_Anexo_Company_Care.docx"
doc.save(output_path)
print("SUCCESS:", output_path)
