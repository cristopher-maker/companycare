import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ==========================================
# 1. GENERATE INTAKE FORM IN WORD (.DOCX)
# ==========================================
doc = docx.Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

COLOR_NAVY = RGBColor(18, 60, 79)       # #123C4F
COLOR_ORANGE = RGBColor(242, 122, 94)   # #F27A5E
COLOR_DARK = RGBColor(44, 62, 80)       # #2C3E50
COLOR_MUTED = RGBColor(107, 123, 133)   # #6B7B85

def set_cell_bg(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_header(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = COLOR_NAVY
    return p

def add_fields_table(fields_list):
    table = doc.add_table(rows=len(fields_list), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (label, default_val) in enumerate(fields_list):
        cell_l = table.rows[idx].cells[0]
        cell_v = table.rows[idx].cells[1]
        
        cell_l.width = Inches(2.5)
        cell_v.width = Inches(4.3)
        
        set_cell_bg(cell_l, "F4F7F9")
        
        p_l = cell_l.paragraphs[0]
        r_l = p_l.add_run(label)
        r_l.font.size = Pt(9.5)
        r_l.font.bold = True
        r_l.font.color.rgb = COLOR_NAVY
        
        p_v = cell_v.paragraphs[0]
        r_v = p_v.add_run(default_val if default_val else "__________________________________________________")
        r_v.font.size = Pt(9.5)
        r_v.font.color.rgb = COLOR_MUTED if not default_val else COLOR_DARK

# Header
p_t = doc.add_paragraph()
r_b = p_t.add_run("CompanyCare | ")
r_b.font.size = Pt(20)
r_b.font.bold = True
r_b.font.color.rgb = COLOR_NAVY

r_title = p_t.add_run("Ficha de Ingreso y Evaluación de Caso")
r_title.font.size = Pt(18)
r_title.font.bold = True
r_title.font.color.rgb = COLOR_ORANGE

p_s = doc.add_paragraph()
r_sub = p_s.add_run("Formulario Oficial de Diagnóstico Familiar y Atención Gerontológica")
r_sub.font.size = Pt(10.5)
r_sub.font.italic = True
r_sub.font.color.rgb = COLOR_MUTED
p_s.paragraph_format.space_after = Pt(10)

# Section 0: Datos del Empleado y Empresa
add_section_header("0. Validación de Beneficio Corporativo (Empleado / Empresa)")
add_fields_table([
    ("Nombre del Colaborador:", ""),
    ("RUT del Colaborador:", ""),
    ("Empresa a la que pertenece:", ""),
    ("Cargo / Área de desempeño:", ""),
    ("Correo corporativo:", ""),
    ("Teléfono de contacto directo:", "")
])

# Section 1: Datos del Adulto Mayor
add_section_header("1. Datos del Adulto Mayor (Beneficiario)")
add_fields_table([
    ("Nombre completo:", ""),
    ("RUT / Documento de Identidad:", ""),
    ("Fecha de Nacimiento:", "DD / MM / AAAA"),
    ("Edad actual:", "______ años"),
    ("Sexo:", "[  ] Femenino   [  ] Masculino   [  ] Otro"),
    ("Nacionalidad:", ""),
    ("Dirección de residencia:", ""),
    ("Comuna / Región:", ""),
    ("Teléfono personal (si aplica):", ""),
    ("Correo electrónico (si aplica):", "")
])

# Section 2: Contacto Responsable
add_section_header("2. Contacto Responsable / Cuidador Principal")
add_fields_table([
    ("Nombre completo del responsable:", ""),
    ("Relación / Parentesco con adulto mayor:", "[  ] Hijo/a   [  ] Cónyuge   [  ] Nieto/a   [  ] Otros: ________"),
    ("RUT / Documento de Identidad:", ""),
    ("Teléfono celular principal:", ""),
    ("Teléfono secundario / Emergencia:", ""),
    ("Correo electrónico:", ""),
    ("Dirección del responsable:", ""),
    ("¿Es el cuidador principal actual?:", "[  ] Sí   [  ] No   (Si es no, indicar quién: _________________)")
])

# Section 3: Información de Salud
add_section_header("3. Información Médica y de Salud")
add_fields_table([
    ("Previsión de salud:", "[  ] Fonasa (Tramo: ___)   [  ] Isapre: _________   [  ] Dipreca/Capredena   [  ] Particular"),
    ("Diagnósticos principales:", ""),
    ("Enfermedades crónicas:", ""),
    ("Medicamentos actuales (dosis/horarios):", ""),
    ("Alergias conocidas:", ""),
    ("Médico tratante / Especialista:", ""),
    ("Teléfono médico / Centro de Salud:", ""),
    ("Última hospitalización (fecha y causa):", ""),
    ("Inscrito en RSH / Carnet Discapacidad:", "[  ] Registro Social de Hogares   [  ] Credencial Discapacidad (Grado: __%)"),
    ("Observaciones médicas relevantes:", "")
])

# Section 4: Nivel de Dependencia
add_section_header("4. Nivel de Dependencia y Funcionalidad Física")
add_fields_table([
    ("Movilidad física:", "[  ] Se moviliza solo/a   [  ] Con ayuda humana   [  ] Silla de ruedas   [  ] Encamado/a"),
    ("Ayudas técnicas en uso:", "[  ] Ninguna   [  ] Bastón   [  ] Andador   [  ] Silla de ruedas   [  ] Catre clínico"),
    ("Ayuda para baño e higiene personal:", "[  ] Independiente   [  ] Requiere ayuda parcial   [  ] Dependencia total"),
    ("Ayuda para vestirse:", "[  ] Independiente   [  ] Requiere ayuda parcial   [  ] Dependencia total"),
    ("Ayuda para alimentarse:", "[  ] Independiente   [  ] Requiere ayuda parcial   [  ] Dependencia total"),
    ("Administración de medicamentos:", "[  ] Autónomo   [  ] Requiere recordatorio   [  ] Requiere administración total"),
    ("Uso de baño / Incontinencia:", "[  ] Autónomo   [  ] Pañal día   [  ] Pañal noche   [  ] Pañal 24 hrs"),
    ("Supervisión requerida:", "[  ] No requiere   [  ] Parcial (Horas)   [  ] Permanente (24/7)"),
    ("Riesgo de caídas:", "[  ] Bajo   [  ] Medio   [  ] Alto   (Antecedentes de caídas recientes: [  ] Sí  [  ] No)")
])

# Section 5: Estado Cognitivo y Emocional
add_section_header("5. Estado Cognitivo, Conductual y Emocional")
add_fields_table([
    ("Orientación espacial y temporal:", "[  ] Orientado/a   [  ] Desorientación ocasional   [  ] Desorientado/a"),
    ("Diagnóstico cognitivo oficial:", "[  ] Sin diagnóstico   [  ] Deterioro leve   [  ] Alzheimer   [  ] Demencia   [  ] Otro: _____"),
    ("Alteraciones de conducta:", "[  ] Ninguna   [  ] Agitación   [  ] Deambulación   [  ] Agresividad   [  ] Trastorno del sueño"),
    ("Estado de ánimo predominante:", "[  ] Estable / Tranquilo   [  ] Tristeza / Aislamiento   [  ] Ansiedad / Angustia"),
    ("Observaciones emocionales:", "")
])

# Section 6: Situación Familiar y Entorno
add_section_header("6. Situación Familiar, Vivienda y Entorno")
add_fields_table([
    ("Composición de la vivienda:", "[  ] Vive solo/a   [  ] Con cónyuge   [  ] Con hijos/familia   [  ] Con cuidador/a"),
    ("Características de la vivienda:", "[  ] Casa 1 piso   [  ] Casa 2 pisos   [  ] Depto con ascensor   [  ] Depto por escalera"),
    ("Adaptaciones en el hogar:", "[  ] Baño adaptado / Barras   [  ] Rampas   [  ] Ninguna adaptativa"),
    ("Red de apoyo familiar disponible:", ""),
    ("Horarios críticos que requieren apoyo:", "[  ] Mañana   [  ] Tarde   [  ] Noche   [  ] Fines de semana   [  ] 24/7"),
    ("Nivel de sobrecarga del cuidador:", "[  ] Leve   [  ] Moderado   [  ] Alto / Sombra de Burnout"),
    ("Urgencia en la resolución del caso:", "[  ] Baja   [  ] Media   [  ] Alta   [  ] Emergencia inmediata")
])

# Section 7: Preferencias y Presupuesto
add_section_header("7. Preferencia de Servicio y Presupuesto Estimado")
add_fields_table([
    ("Modalidad de cuidado buscada:", "[  ] Por horas   [  ] Turno día   [  ] Turno noche   [  ] Cuidador interno (24/7)   [  ] Residencia"),
    ("Presupuesto mensual aproximado:", "[  ] $400.000 - $700.000   [  ] $700.000 - $1.200.000   [  ] Sobre $1.200.000"),
    ("Comuna / Sector preferido:", ""),
    ("Fecha estimada de inicio del servicio:", "DD / MM / AAAA")
])

# Section 8: Servicio Solicitado
add_section_header("8. Tipo de Servicio Requerido (Marcar opciones)")
add_fields_table([
    ("Servicios requeridos (marcar todos):", 
     "[  ] Asesoría con Care Experts\n"
     "[  ] Búsqueda de Cuidador/a Domiciliario/a\n"
     "[  ] Selección de Residencia (ELEAM)\n"
     "[  ] Servicio de Enfermería / Kinesiología\n"
     "[  ] Apoyo Psicológico al Cuidador\n"
     "[  ] Orientación Legal y Trámites Previsionales\n"
     "[  ] Tramitación de Bonos y Subsidios Senior\n"
     "[  ] Otro: ____________________________________")
])

# Section 9: Descripción del Caso y Firma
add_section_header("9. Descripción del Caso y Consentimiento")
p_desc = doc.add_paragraph()
r_d1 = p_desc.add_run("Describa brevemente la situación actual, los problemas principales a resolver y qué esperan de CompanyCare:\n\n")
r_d1.font.size = Pt(9.5)
r_d1.font.bold = True
r_d1.font.color.rgb = COLOR_NAVY

r_lines = p_desc.add_run(
    "____________________________________________________________________________________________________\n"
    "____________________________________________________________________________________________________\n"
    "____________________________________________________________________________________________________\n\n"
)
r_lines.font.size = Pt(9)
r_lines.font.color.rgb = COLOR_MUTED

p_firm = doc.add_paragraph()
p_firm.paragraph_format.space_before = Pt(16)
r_f = p_firm.add_run(
    "Declaro que la información entregada es fidedigna y autorizo el tratamiento de datos para la gestión del caso.\n\n\n"
    "_________________________________________                  _________________________\n"
    "   Firma del Empleado / Responsable                                    Fecha"
)
r_f.font.size = Pt(9)
r_f.font.color.rgb = COLOR_DARK

doc.save("c:/Users/crist/OneDrive/Documentos/Company Care/Ficha_Ingreso_Caso_CompanyCare.docx")
print("Intake Word generated successfully.")

# ==========================================
# 2. GENERATE INTAKE FORM IN EXCEL (.XLSX)
# ==========================================
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Ficha Ingreso Caso Senior"
ws.views.sheetView[0].showGridLines = True

font_t = Font(name="Segoe UI", size=15, bold=True, color="123C4F")
font_s = Font(name="Segoe UI", size=10, italic=True, color="6B7B85")
font_sec = Font(name="Segoe UI", size=11, bold=True, color="FFFFFF")
font_lbl = Font(name="Segoe UI", size=9.5, bold=True, color="123C4F")
font_val = Font(name="Segoe UI", size=9.5, color="2C3E50")

fill_sec = PatternFill(start_color="123C4F", end_color="123C4F", fill_type="solid")
fill_lbl = PatternFill(start_color="F4F7F9", end_color="F4F7F9", fill_type="solid")

thin_b = Border(
    left=Side(style='thin', color='E5E9EB'),
    right=Side(style='thin', color='E5E9EB'),
    top=Side(style='thin', color='E5E9EB'),
    bottom=Side(style='thin', color='E5E9EB')
)

ws["A1"] = "CompanyCare - Ficha de Ingreso y Evaluación de Caso"
ws["A1"].font = font_t
ws["A2"] = "Formulario Oficial de Diagnóstico Familiar y Atención Gerontológica"
ws["A2"].font = font_s

intake_sections = [
    ("0. Datos del Colaborador y Empresa", [
        ("Nombre Colaborador", ""), ("RUT Colaborador", ""), ("Empresa Beneficiaria", ""), 
        ("Cargo / Área", ""), ("Correo Corporativo", ""), ("Teléfono", "")
    ]),
    ("1. Datos del Adulto Mayor", [
        ("Nombre Completo", ""), ("RUT / Doc", ""), ("Fecha Nacimiento", "DD/MM/AAAA"), 
        ("Edad", ""), ("Sexo", "Femenino / Masculino / Otro"), ("Nacionalidad", ""),
        ("Dirección", ""), ("Comuna / Región", ""), ("Teléfono", ""), ("Correo", "")
    ]),
    ("2. Contacto Responsable", [
        ("Nombre Responsable", ""), ("Parentesco", "Hijo/a / Cónyuge / Nieto/a / Otro"),
        ("RUT Responsable", ""), ("Teléfono Celular", ""), ("Teléfono Emergencia", ""),
        ("Correo", ""), ("Dirección", ""), ("¿Cuidador Principal?", "Sí / No")
    ]),
    ("3. Información de Salud", [
        ("Previsión Salud", "Fonasa / Isapre / Dipreca / Particular"), ("Diagnósticos Principales", ""),
        ("Enfermedades Crónicas", ""), ("Medicamentos (Dosis)", ""), ("Alergias", ""),
        ("Médico Tratante", ""), ("Teléfono Médico/Centro", ""), ("Última Hospitalización", ""),
        ("RSH / Credencial Discapacidad", "Registro Social Hogares / Carnet Discapacidad")
    ]),
    ("4. Nivel de Dependencia Física", [
        ("Movilidad", "Solo/a / Con ayuda / Silla de ruedas / Encamado"), 
        ("Ayudas Técnicas", "Bastón / Andador / Silla / Catre clínico"),
        ("Ayuda en Baño/Higiene", "Independiente / Ayuda Parcial / Dependencia Total"),
        ("Ayuda para Vestirse", "Independiente / Ayuda Parcial / Dependencia Total"),
        ("Ayuda Alimentación", "Independiente / Ayuda Parcial / Dependencia Total"),
        ("Toma Medicamentos", "Autónomo / Recordatorio / Administración Total"),
        ("Uso Baño / Incontinencia", "Autónomo / Pañal Día / Pañal Noche / Pañal 24h"),
        ("Supervisión Requerida", "No requiere / Parcial / Permanente 24/7"),
        ("Riesgo Caídas", "Bajo / Medio / Alto")
    ]),
    ("5. Estado Cognitivo y Emocional", [
        ("Orientación Espacio/Tiempo", "Orientado/a / Parcial / Desorientado/a"),
        ("Diagnóstico Cognitivo", "Sin diag. / Deterioro Leve / Alzheimer / Demencia"),
        ("Cambios de Conducta", "Agitación / Deambulación / Agresividad / Sueño"),
        ("Estado de Ánimo", "Estable / Tristeza / Aislamiento / Ansiedad")
    ]),
    ("6. Situación Familiar y Entorno", [
        ("Composición Vivienda", "Solo/a / Cónyuge / Hijos / Cuidador"),
        ("Tipo Vivienda", "Casa 1 piso / Casa 2 pisos / Depto ascensor / Depto escalera"),
        ("Horarios Apoyo Requerido", "Mañana / Tarde / Noche / Fines de semana / 24/7"),
        ("Sobrecarga Cuidador", "Leve / Moderada / Alta / Burnout"),
        ("Urgencia del Caso", "Baja / Media / Alta / Emergencia")
    ]),
    ("7. Servicio Solicitado y Presupuesto", [
        ("Modalidad Cuidado", "Por horas / Turno día / Turno noche / Interno 24/7 / Residencia"),
        ("Presupuesto Estimado", "$400k-$700k / $700k-$1.2M / Sobre $1.2M"),
        ("Comuna / Sector Preferido", ""),
        ("Servicios Requeridos", "Asesoría / Cuidador / Residencia / Enfermería / Psicología / Legal")
    ]),
    ("8. Descripción del Caso", [
        ("Resumen Situación Actual", "Escriba aquí los detalles del caso y necesidades principales...")
    ])
]

curr_r = 4
for sec_title, fields in intake_sections:
    # Section Header Row
    ws.cell(row=curr_r, column=1, value=sec_title)
    ws.cell(row=curr_r, column=2, value="")
    ws.merge_cells(start_row=curr_r, start_column=1, end_row=curr_r, end_column=2)
    
    cell_h = ws.cell(row=curr_r, column=1)
    cell_h.font = font_sec
    cell_h.fill = fill_sec
    cell_h.alignment = Alignment(horizontal="left", vertical="center")
    curr_r += 1
    
    for f_lbl, f_val in fields:
        c1 = ws.cell(row=curr_r, column=1, value=f_lbl)
        c2 = ws.cell(row=curr_r, column=2, value=f_val)
        
        c1.font = font_lbl
        c1.fill = fill_lbl
        c1.border = thin_b
        
        c2.font = font_val
        c2.border = thin_b
        curr_r += 1
    curr_r += 1

# Auto adjust columns
ws.column_dimensions["A"].width = 34
ws.column_dimensions["B"].width = 60

wb.save("c:/Users/crist/OneDrive/Documentos/Company Care/Ficha_Ingreso_Caso_CompanyCare.xlsx")
print("Intake Excel generated successfully.")
